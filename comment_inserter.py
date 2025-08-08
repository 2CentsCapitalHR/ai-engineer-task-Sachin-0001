import os
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class CommentInserter:
    """Inserts contextual comments into .docx documents based on analysis results."""
    
    def __init__(self):
        self.comment_counter = 0
    
    def add_comment_to_paragraph(self, paragraph, comment_text: str, author: str = "ADGM Agent"):
        """Add a comment to a specific paragraph."""
        self.comment_counter += 1
        
        # Create comment element
        comment_element = OxmlElement('w:commentRangeStart')
        comment_element.set(qn('w:id'), str(self.comment_counter))
        
        comment_end_element = OxmlElement('w:commentRangeEnd')
        comment_end_element.set(qn('w:id'), str(self.comment_counter))
        
        comment_reference = OxmlElement('w:commentReference')
        comment_reference.set(qn('w:id'), str(self.comment_counter))
        
        # Insert comment elements
        paragraph._element.insert(0, comment_element)
        paragraph._element.append(comment_end_element)
        paragraph._element.append(comment_reference)
        
        # Add comment to document
        self._add_comment_to_document(comment_text, author, str(self.comment_counter))
    
    def _add_comment_to_document(self, comment_text: str, author: str, comment_id: str):
        """Add comment to document's comment collection."""
        # This is a simplified version - in a full implementation,
        # you would need to properly add to the document's comment collection
        pass
    
    def insert_comments_based_on_issues(self, doc: Document, issues: List[Dict[str, Any]], 
                                      document_text: str) -> Document:
        """Insert comments into document based on identified issues."""
        if not issues:
            return doc
        
        # Create a mapping of issue types to comment templates
        comment_templates = {
            "jurisdiction_issue": {
                "template": "🚨 JURISDICTION ISSUE: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF0000"  # Red
            },
            "missing_clause": {
                "template": "⚠️ MISSING CLAUSE: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF6600"  # Orange
            },
            "ambiguous_language": {
                "template": "⚠️ AMBIGUOUS LANGUAGE: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF9900"  # Yellow
            },
            "missing_signatures": {
                "template": "🚨 MISSING SIGNATURES: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF0000"  # Red
            },
            "incomplete_info": {
                "template": "⚠️ INCOMPLETE INFORMATION: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF9900"  # Yellow
            },
            "non_compliant_structure": {
                "template": "🚨 NON-COMPLIANT STRUCTURE: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "FF0000"  # Red
            },
            "formatting_issue": {
                "template": "ℹ️ FORMATTING ISSUE: {description}\n\nADGM Reference: {adgm_reference}\n\nSuggestion: {suggestion}",
                "severity_color": "0066CC"  # Blue
            }
        }
        
        # Add summary comment at the beginning
        summary_comment = self._create_summary_comment(issues)
        if doc.paragraphs:
            self.add_comment_to_paragraph(doc.paragraphs[0], summary_comment, "ADGM Agent")
        
        # Add specific comments for each issue
        for issue in issues:
            issue_type = issue.get("type", "unknown")
            if issue_type in comment_templates:
                template = comment_templates[issue_type]["template"]
                comment_text = template.format(
                    description=issue.get("description", ""),
                    adgm_reference=issue.get("adgm_reference", "ADGM Companies Regulations 2020"),
                    suggestion=issue.get("suggestion", "")
                )
                
                # Find appropriate paragraph to comment on
                target_paragraph = self._find_target_paragraph(doc, issue, document_text)
                if target_paragraph:
                    self.add_comment_to_paragraph(target_paragraph, comment_text, "ADGM Agent")
        
        return doc
    
    def _create_summary_comment(self, issues: List[Dict[str, Any]]) -> str:
        """Create a summary comment for all issues."""
        if not issues:
            return "✅ No issues found. Document appears to be compliant with ADGM requirements."
        
        # Count issues by severity
        severity_counts = {"High": 0, "Medium": 0, "Low": 0}
        for issue in issues:
            severity = issue.get("severity", "Medium")
            severity_counts[severity] += 1
        
        summary = f"""📋 ADGM COMPLIANCE ANALYSIS SUMMARY

Total Issues Found: {len(issues)}
- High Severity: {severity_counts['High']}
- Medium Severity: {severity_counts['Medium']}
- Low Severity: {severity_counts['Low']}

Please review all comments throughout the document for detailed recommendations.
"""
        
        return summary
    
    def _find_target_paragraph(self, doc: Document, issue: Dict[str, Any], 
                             document_text: str) -> Any:
        """Find the most appropriate paragraph to attach a comment to."""
        # For now, return the first paragraph with content
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                return paragraph
        
        return None
    
    def add_highlighted_sections(self, doc: Document, issues: List[Dict[str, Any]]) -> Document:
        """Add highlighting to problematic sections."""
        # This would require more sophisticated text matching
        # For now, we'll add a general highlighting approach
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower()
            
            # Highlight jurisdiction issues
            if any(keyword in text for keyword in ["uae federal", "federal courts", "uae law"]):
                paragraph.style = 'Intense Quote'
            
            # Highlight missing signature indicators
            if any(keyword in text for keyword in ["signature", "signed", "executed"]):
                if not any(keyword in text for keyword in ["signature block", "signed by", "executed by"]):
                    paragraph.style = 'Intense Quote'
        
        return doc
    
    def create_reviewed_document(self, original_path: str, issues: List[Dict[str, Any]], 
                               output_path: str) -> str:
        """Create a reviewed version of the document with comments."""
        try:
            # Load original document
            doc = Document(original_path)
            
            # Extract text for analysis
            document_text = ""
            for paragraph in doc.paragraphs:
                document_text += paragraph.text + "\n"
            
            # Insert comments based on issues
            doc = self.insert_comments_based_on_issues(doc, issues, document_text)
            
            # Add highlighting
            doc = self.add_highlighted_sections(doc, issues)
            
            # Save reviewed document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating reviewed document: {str(e)}")
    
    def generate_comment_text(self, issue: Dict[str, Any]) -> str:
        """Generate formatted comment text for an issue."""
        severity_icons = {
            "High": "🚨",
            "Medium": "⚠️", 
            "Low": "ℹ️"
        }
        
        icon = severity_icons.get(issue.get("severity", "Medium"), "ℹ️")
        
        comment = f"""{icon} {issue.get('type', 'ISSUE').replace('_', ' ').upper()}

{issue.get('description', 'No description provided')}

ADGM Reference: {issue.get('adgm_reference', 'ADGM Companies Regulations 2020')}

Suggestion: {issue.get('suggestion', 'Please review and correct.')}
"""
        
        return comment
    
    def add_general_compliance_notes(self, doc: Document, document_type: str) -> Document:
        """Add general compliance notes based on document type."""
        compliance_notes = {
            "Articles of Association": """
✅ ADGM COMPLIANCE CHECKLIST:
- Ensure ADGM jurisdiction is specified
- Include complete objects clause
- Specify share capital structure
- Include director appointment procedures
- Add shareholder rights and meetings
- Include registered office details
""",
            "Memorandum of Association": """
✅ ADGM COMPLIANCE CHECKLIST:
- Include company name and type
- Specify registered office address
- Include objects and powers
- Add liability of members
- Include capital structure
""",
            "Board Resolution": """
✅ ADGM COMPLIANCE CHECKLIST:
- Include meeting date and location
- List all directors present
- Include complete resolution text
- Add voting results
- Include signature blocks
""",
            "UBO Declaration": """
✅ ADGM COMPLIANCE CHECKLIST:
- Include all beneficial owner details
- Specify ownership percentages
- Include control structures
- Add declaration statements
- Include supporting documentation
"""
        }
        
        if document_type in compliance_notes:
            # Add compliance note as first paragraph
            note_paragraph = doc.add_paragraph()
            note_paragraph.text = compliance_notes[document_type]
            note_paragraph.style = 'Quote'
        
        return doc
